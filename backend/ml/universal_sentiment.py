# file: universal_sentiment.py
"""
Universal Sentiment Analyzer - Handles PDF, DOC, CSV, Excel, TXT, JSON files
"""

import os
import re
import json
import pandas as pd
import numpy as np
from datetime import datetime
from typing import Dict, List, Any, Optional
import pickle
import warnings
warnings.filterwarnings('ignore')

# Try to import optional libraries
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️  PyPDF2 not installed. PDF support limited.")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️  python-docx not installed. DOCX support limited.")

class FileProcessor:
    """Process multiple file formats for sentiment analysis"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.json': self._process_json,
        }
        
        # Add optional formats if libraries are available
        if PDF_SUPPORT:
            self.supported_formats['.pdf'] = self._process_pdf
        
        if DOCX_SUPPORT:
            self.supported_formats['.docx'] = self._process_docx
        
        # Try to add .doc support
        try:
            import textract
            self.supported_formats['.doc'] = self._process_doc
        except:
            pass
    
    def process_file(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process any supported file and extract text"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            # Try to read as text file
            return self._process_unknown(file_path)
        
        try:
            result = self.supported_formats[file_ext](file_path, text_column)
            return result
        except Exception as e:
            print(f"Error processing {file_ext} file: {e}")
            return self._process_unknown(file_path)
    
    def _process_txt(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content) if p.strip()]
        
        return {
            'file_type': 'txt',
            'file_name': os.path.basename(file_path),
            'texts': paragraphs if paragraphs else [content],
            'raw_text': content,
            'total_texts': len(paragraphs) if paragraphs else 1
        }
    
    def _process_csv(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except:
            try:
                df = pd.read_csv(file_path, encoding='latin-1')
            except:
                df = pd.read_csv(file_path, encoding='utf-16')
        
        texts = self._extract_texts_from_dataframe(df, text_column)
        
        return {
            'file_type': 'csv',
            'file_name': os.path.basename(file_path),
            'texts': texts,
            'dataframe': df,
            'total_texts': len(texts)
        }
    
    def _process_excel(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            df = pd.read_excel(file_path)
        except Exception as e:
            print(f"Excel read error: {e}")
            df = pd.read_excel(file_path, sheet_name=0)
        
        texts = self._extract_texts_from_dataframe(df, text_column)
        
        return {
            'file_type': 'excel',
            'file_name': os.path.basename(file_path),
            'texts': texts,
            'dataframe': df,
            'total_texts': len(texts)
        }
    
    def _process_pdf(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process PDF file"""
        if not PDF_SUPPORT:
            return self._process_unknown(file_path)
        
        try:
            reader = PdfReader(file_path)
            text_content = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(text)
            
            full_text = "\n\n".join(text_content)
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', full_text) if p.strip()]
            
            return {
                'file_type': 'pdf',
                'file_name': os.path.basename(file_path),
                'texts': paragraphs if paragraphs else text_content,
                'raw_text': full_text,
                'total_pages': len(reader.pages),
                'total_texts': len(paragraphs) if paragraphs else len(text_content)
            }
            
        except Exception as e:
            print(f"PDF processing error: {e}")
            return self._process_unknown(file_path)
    
    def _process_docx(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process DOCX file"""
        if not DOCX_SUPPORT:
            return self._process_unknown(file_path)
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        table_texts.append(row_text)
            
            all_texts = paragraphs + table_texts
            full_text = "\n\n".join(all_texts)
            
            return {
                'file_type': 'docx',
                'file_name': os.path.basename(file_path),
                'texts': all_texts,
                'raw_text': full_text,
                'total_texts': len(all_texts)
            }
        except Exception as e:
            print(f"DOCX processing error: {e}")
            return self._process_unknown(file_path)
    
    def _process_doc(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process DOC file"""
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            
            paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
            
            return {
                'file_type': 'doc',
                'file_name': os.path.basename(file_path),
                'texts': paragraphs,
                'raw_text': text,
                'total_texts': len(paragraphs)
            }
        except Exception as e:
            print(f"DOC processing error: {e}")
            return self._process_unknown(file_path)
    
    def _process_json(self, file_path: str, text_column: str = None) -> Dict[str, Any]:
        """Process JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        texts = self._extract_texts_from_json(data, text_column)
        
        return {
            'file_type': 'json',
            'file_name': os.path.basename(file_path),
            'texts': texts,
            'total_texts': len(texts)
        }
    
    def _process_unknown(self, file_path: str) -> Dict[str, Any]:
        """Process unknown file as text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        except:
            with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
                content = f.read()
        
        paragraphs = [p.strip() for p in re.split(r'\n\s*\n', content) if p.strip()]
        
        return {
            'file_type': 'unknown',
            'file_name': os.path.basename(file_path),
            'texts': paragraphs if paragraphs else [content],
            'raw_text': content,
            'total_texts': len(paragraphs) if paragraphs else 1
        }
    
    def _extract_texts_from_dataframe(self, df: pd.DataFrame, text_column: str = None) -> List[str]:
        """Extract texts from DataFrame"""
        if text_column and text_column in df.columns:
            texts = df[text_column].dropna().astype(str).tolist()
        else:
            # Auto-detect text column
            text_columns = [col for col in df.columns if df[col].dtype == 'object']
            if text_columns:
                # Try to find column with longest average text
                col_lengths = []
                for col in text_columns:
                    avg_len = df[col].astype(str).str.len().mean()
                    col_lengths.append((col, avg_len))
                
                if col_lengths:
                    best_col = max(col_lengths, key=lambda x: x[1])[0]
                    texts = df[best_col].dropna().astype(str).tolist()
                else:
                    texts = df.astype(str).apply(lambda row: ' '.join(row), axis=1).tolist()
            else:
                texts = df.astype(str).apply(lambda row: ' '.join(row), axis=1).tolist()
        
        # Clean texts
        texts = [self._clean_text(text) for text in texts if str(text).strip()]
        return texts
    
    def _extract_texts_from_json(self, data: Any, text_column: str = None) -> List[str]:
        """Extract texts from JSON structure"""
        texts = []
        
        if isinstance(data, dict):
            if text_column and text_column in data:
                texts.append(str(data[text_column]))
            else:
                # Extract all string values
                for value in data.values():
                    if isinstance(value, str) and len(value.strip()) > 10:
                        texts.append(value)
                    elif isinstance(value, (dict, list)):
                        texts.extend(self._extract_texts_from_json(value, text_column))
        elif isinstance(data, list):
            for item in data:
                texts.extend(self._extract_texts_from_json(item, text_column))
        
        # Clean and deduplicate
        texts = [self._clean_text(text) for text in texts if str(text).strip()]
        return list(set(texts))
    
    def _clean_text(self, text: str) -> str:
        """Clean text"""
        text = str(text)
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove non-printable characters
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        return text.strip()

class SentimentAnalyzer:
    """Sentiment analysis with fallback models"""
    
    def __init__(self):
        self.model = None
        self.vectorizer = None
        self.model_source = "unknown"
        self._load_best_model()
    
    def _load_best_model(self):
        """Load the best available model"""
        model_paths = [
            ('expert', 'expert_sentiment_model.pkl', 'expert_vectorizer.pkl'),
            ('best', 'best_sentiment_model.pkl', 'best_vectorizer.pkl'),
            ('improved', 'improved_working_model.pkl', 'improved_working_vectorizer.pkl')
        ]
        
        for source, model_file, vectorizer_file in model_paths:
            try:
                if os.path.exists(model_file) and os.path.exists(vectorizer_file):
                    with open(model_file, 'rb') as f:
                        self.model = pickle.load(f)
                    with open(vectorizer_file, 'rb') as f:
                        self.vectorizer = pickle.load(f)
                    self.model_source = source
                    print(f"✅ Loaded {source} sentiment model")
                    return True
            except Exception as e:
                print(f"⚠️  Failed to load {source} model: {e}")
        
        # Create fallback model
        print("🔄 Creating fallback model...")
        self._create_fallback_model()
        return True
    
    def _create_fallback_model(self):
        """Create a simple fallback model"""
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.linear_model import LogisticRegression
        
        # Simple training data
        texts = [
            "good great excellent love happy",
            "bad terrible awful hate angry",
            "okay fine average normal neutral"
        ]
        labels = [1, 0, 2]  # 1=positive, 0=negative, 2=neutral
        
        self.vectorizer = TfidfVectorizer()
        X = self.vectorizer.fit_transform(texts)
        
        self.model = LogisticRegression(solver='lbfgs', max_iter=1000)
        self.model.fit(X, labels)
        self.model_source = "fallback"
    
    def analyze_sentiment(self, text: str) -> Dict[str, Any]:
        """Analyze sentiment of a single text"""
        if not self.model or not self.vectorizer:
            return {"error": "Models not loaded"}
        
        try:
            X = self.vectorizer.transform([text])
            prediction = self.model.predict(X)[0]
            
            if hasattr(self.model, 'predict_proba'):
                probabilities = self.model.predict_proba(X)[0]
                confidence = float(np.max(probabilities))
                
                # Create probability dict
                sentiment_map = {0: "negative", 1: "positive", 2: "neutral"}
                prob_dict = {}
                for i, class_label in enumerate(self.model.classes_):
                    label_name = sentiment_map.get(class_label, f"class_{class_label}")
                    prob_dict[label_name] = float(probabilities[i])
            else:
                confidence = 1.0
                prob_dict = {}
            
            sentiment_map = {0: "negative", 1: "positive", 2: "neutral"}
            sentiment = sentiment_map.get(prediction, "unknown")
            
            # Calculate sentiment score (-1 to 1)
            sentiment_score = self._calculate_sentiment_score(prob_dict)
            
            return {
                "success": True,
                "text": text,
                "sentiment": sentiment,
                "confidence": confidence,
                "sentiment_score": sentiment_score,
                "probabilities": prob_dict,
                "metadata": {
                    "model_source": self.model_source,
                    "processed_at": datetime.now().isoformat(),
                    "text_length": len(text),
                    "word_count": len(text.split())
                }
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": text
            }
    
    def _calculate_sentiment_score(self, probabilities: Dict) -> float:
        """Calculate sentiment score from -1 to 1"""
        neg = probabilities.get('negative', 0)
        pos = probabilities.get('positive', 0)
        score = pos - neg
        return round(score, 3)
    
    def analyze_batch(self, texts: List[str]) -> Dict[str, Any]:
        """Analyze multiple texts"""
        results = []
        sentiment_counts = {"positive": 0, "negative": 0, "neutral": 0}
        total_confidence = 0
        total_score = 0
        
        for text in texts:
            if not text or not isinstance(text, str):
                continue
                
            result = self.analyze_sentiment(text.strip())
            if result.get("success"):
                results.append(result)
                sentiment = result["sentiment"]
                if sentiment in sentiment_counts:
                    sentiment_counts[sentiment] += 1
                total_confidence += result["confidence"]
                total_score += result["sentiment_score"]
        
        total_results = len(results)
        
        if total_results > 0:
            avg_confidence = total_confidence / total_results
            avg_score = total_score / total_results
            
            sentiment_percentages = {
                sentiment: count / total_results * 100 
                for sentiment, count in sentiment_counts.items()
            }
            
            dominant = max(sentiment_counts.items(), key=lambda x: x[1])[0]
            
            return {
                "success": True,
                "total_texts": len(texts),
                "analyzed_texts": total_results,
                "sentiment_counts": sentiment_counts,
                "sentiment_percentages": sentiment_percentages,
                "dominant_sentiment": dominant,
                "average_confidence": avg_confidence,
                "average_sentiment_score": avg_score,
                "model_source": self.model_source,
                "results": results[:100]  # Limit to first 100 for performance
            }
        else:
            return {
                "success": False,
                "error": "No texts could be analyzed",
                "total_texts": len(texts),
                "analyzed_texts": 0
            }

class UniversalSentimentAnalyzer:
    """
    Main class for universal sentiment analysis.
    Usage:
        analyzer = UniversalSentimentAnalyzer()
        results = analyzer.analyze_file("document.pdf")
        results = analyzer.analyze_text("I love this product!")
    """
    
    def __init__(self):
        self.file_processor = FileProcessor()
        self.sentiment_analyzer = SentimentAnalyzer()
    
    def analyze_file(self, file_path: str, text_column: str = None, 
                    max_texts: int = 1000) -> Dict[str, Any]:
        """
        Analyze sentiment from a file.
        
        Args:
            file_path: Path to the file (PDF, DOC, CSV, Excel, TXT, JSON)
            text_column: For structured files, column containing text
            max_texts: Maximum number of texts to process (for performance)
        
        Returns:
            Dictionary with analysis results
        """
        print(f"📄 Processing file: {os.path.basename(file_path)}")
        
        try:
            # Process file
            file_data = self.file_processor.process_file(file_path, text_column)
            
            if not file_data.get('texts'):
                return {
                    "success": False,
                    "error": "No text content found in file",
                    "file_info": file_data
                }
            
            print(f"   Found {len(file_data['texts'])} text segments")
            
            # Limit texts for performance
            texts_to_analyze = file_data['texts'][:max_texts]
            
            # Analyze sentiment
            sentiment_results = self.sentiment_analyzer.analyze_batch(texts_to_analyze)
            
            # Combine results
            results = {
                "success": True,
                "file_info": {
                    "file_name": file_data['file_name'],
                    "file_type": file_data['file_type'],
                    "total_texts": len(file_data['texts']),
                    "processed_texts": len(texts_to_analyze)
                },
                "analysis": sentiment_results,
                "processed_at": datetime.now().isoformat()
            }
            
            return results
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }
    
    def analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Analyze sentiment from direct text.
        
        Args:
            text: Text to analyze
        
        Returns:
            Dictionary with analysis results
        """
        result = self.sentiment_analyzer.analyze_sentiment(text)
        
        return {
            "success": result.get("success", False),
            "input_type": "direct_text",
            "analysis": result,
            "processed_at": datetime.now().isoformat()
        }
    
    def analyze_directory(self, dir_path: str, text_column: str = None,
                         max_files: int = 50, max_texts_per_file: int = 100) -> Dict[str, Any]:
        """
        Analyze all supported files in a directory.
        
        Args:
            dir_path: Path to directory
            text_column: For structured files, column containing text
            max_files: Maximum number of files to process
            max_texts_per_file: Maximum texts per file
        
        Returns:
            Dictionary with analysis results
        """
        print(f"📁 Processing directory: {dir_path}")
        
        supported_extensions = list(self.file_processor.supported_formats.keys())
        
        # Find all supported files
        all_files = []
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in supported_extensions:
                    all_files.append(os.path.join(root, file))
        
        print(f"   Found {len(all_files)} supported files")
        
        # Process files
        file_results = []
        for i, file_path in enumerate(all_files[:max_files]):
            print(f"   [{i+1}/{min(len(all_files), max_files)}] {os.path.basename(file_path)}")
            
            try:
                file_result = self.analyze_file(file_path, text_column, max_texts_per_file)
                file_results.append(file_result)
            except Exception as e:
                print(f"      ❌ Error: {e}")
                file_results.append({
                    "success": False,
                    "error": str(e),
                    "file_path": file_path
                })
        
        # Compile summary
        successful_files = [f for f in file_results if f.get("success")]
        
        if successful_files:
            # Aggregate sentiment across all files
            total_positive = sum(f.get("analysis", {}).get("sentiment_counts", {}).get("positive", 0) 
                               for f in successful_files)
            total_negative = sum(f.get("analysis", {}).get("sentiment_counts", {}).get("negative", 0) 
                               for f in successful_files)
            total_neutral = sum(f.get("analysis", {}).get("sentiment_counts", {}).get("neutral", 0) 
                              for f in successful_files)
            
            total_sentiments = total_positive + total_negative + total_neutral
            
            if total_sentiments > 0:
                sentiment_percentages = {
                    "positive": total_positive / total_sentiments * 100,
                    "negative": total_negative / total_sentiments * 100,
                    "neutral": total_neutral / total_sentiments * 100
                }
                
                dominant = max(["positive", "negative", "neutral"], 
                             key=lambda x: sentiment_percentages[x])
            else:
                sentiment_percentages = {"positive": 0, "negative": 0, "neutral": 0}
                dominant = "unknown"
            
            return {
                "success": True,
                "directory": dir_path,
                "total_files": len(all_files),
                "processed_files": len(file_results),
                "successful_files": len(successful_files),
                "overall_sentiment": {
                    "counts": {
                        "positive": total_positive,
                        "negative": total_negative,
                        "neutral": total_neutral
                    },
                    "percentages": sentiment_percentages,
                    "dominant": dominant
                },
                "file_results": file_results,
                "processed_at": datetime.now().isoformat()
            }
        else:
            return {
                "success": False,
                "error": "No files could be processed successfully",
                "directory": dir_path,
                "total_files": len(all_files),
                "processed_files": len(file_results)
            }

def analyze_sentiment_from_file(file_path: str, text_column: str = None) -> Dict[str, Any]:
    """
    Convenience function to analyze sentiment from a file.
    
    Args:
        file_path: Path to file
        text_column: Column containing text (for CSV/Excel)
    
    Returns:
        Analysis results as dictionary
    """
    analyzer = UniversalSentimentAnalyzer()
    return analyzer.analyze_file(file_path, text_column)

def analyze_sentiment_from_text(text: str) -> Dict[str, Any]:
    """
    Convenience function to analyze sentiment from text.
    
    Args:
        text: Text to analyze
    
    Returns:
        Analysis results as dictionary
    """
    analyzer = UniversalSentimentAnalyzer()
    return analyzer.analyze_text(text)

if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 2:
        print("Universal Sentiment Analyzer")
        print("=" * 50)
        print("Usage:")
        print("  python universal_sentiment.py <file_path> [text_column]")
        print("  python universal_sentiment.py --text \"Your text here\"")
        print("\nExamples:")
        print("  python universal_sentiment.py document.pdf")
        print("  python universal_sentiment.py reviews.csv --column comments")
        print("  python universal_sentiment.py --text \"I love this product!\"")
        sys.exit(1)
    
    if sys.argv[1] == "--text":
        text = " ".join(sys.argv[2:])
        results = analyze_sentiment_from_text(text)
    else:
        file_path = sys.argv[1]
        text_column = sys.argv[2] if len(sys.argv) > 2 else None
        results = analyze_sentiment_from_file(file_path, text_column)
    
    # Print results
    print("\n📊 Analysis Results:")
    print(json.dumps(results, indent=2, ensure_ascii=False))
